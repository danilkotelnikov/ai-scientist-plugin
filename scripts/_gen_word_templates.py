"""One-shot helper: generate 3 Word manuscript templates.

Run this once to (re)generate mcp/templates/word/*.docx. The plugin's runtime
doesn't depend on this script — it produces static .docx files that ship with
the plugin.
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PLUGIN_ROOT = Path(__file__).resolve().parent.parent
DST = PLUGIN_ROOT / "mcp" / "templates" / "word"
DST.mkdir(parents=True, exist_ok=True)


def build_arxiv():
    d = Document()
    for s in d.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
    style = d.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    d.add_heading("%TITLE%", 0)
    d.add_paragraph("%AUTHOR%")
    d.add_paragraph("%DATE%")
    for sec in ("Abstract", "Introduction", "Methods", "Results", "Discussion", "Conclusion", "References"):
        d.add_heading(sec, 1)
        d.add_paragraph(f"%{sec.upper()}_BODY%")
    d.save(str(DST / "arxiv-shared-1.docx"))


def build_minimalist():
    d = Document()
    for s in d.sections:
        s.top_margin = Inches(1.25)
        s.bottom_margin = Inches(1.25)
        s.left_margin = Inches(1.25)
        s.right_margin = Inches(1.25)
    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    d.add_heading("%TITLE%", 0)
    d.add_paragraph("%AUTHOR%")
    for sec in ("Abstract", "Introduction", "Methods", "Results", "Discussion", "Conclusion", "References"):
        d.add_heading(sec, 1)
        d.add_paragraph(f"%{sec.upper()}_BODY%")
    d.save(str(DST / "minimalist.docx"))


def build_two_column():
    d = Document()
    section = d.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "432")
    style = d.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    d.add_heading("%TITLE%", 0)
    d.add_paragraph("%AUTHOR%")
    for sec in ("Abstract", "Introduction", "Methods", "Results", "Discussion", "Conclusion", "References"):
        d.add_heading(sec, 1)
        d.add_paragraph(f"%{sec.upper()}_BODY%")
    d.save(str(DST / "two-column-academic.docx"))


if __name__ == "__main__":
    build_arxiv()
    print(f"Wrote {DST / 'arxiv-shared-1.docx'}")
    build_minimalist()
    print(f"Wrote {DST / 'minimalist.docx'}")
    build_two_column()
    print(f"Wrote {DST / 'two-column-academic.docx'}")
